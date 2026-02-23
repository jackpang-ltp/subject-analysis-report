import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

# --- Configuration ---
st.set_page_config(page_title="School Report Generator", layout="wide")

# --- Helper Functions ---
def clean_numeric(val):
    if isinstance(val, str):
        val = val.strip()
        if val == '--' or val == '':
            return np.nan
    try:
        return float(val)
    except ValueError:
        return np.nan

def get_sorted_groups(unique_items):
    # Sorts classes or subject blocks
    # Priority order: S, M, A, R, T
    order = {'S': 1, 'M': 2, 'A': 3, 'R': 4, 'T': 5}
    
    def sort_key(name):
        name = str(name)
        # Find the first letter that matches S, M, A, R, T
        match = re.search(r'[SMART]', name)
        if match:
            return order.get(match.group(0), 99)
        return 99 # Put others at the end

    return sorted([x for x in unique_items if pd.notna(x)], key=sort_key)

def create_chart_image(fig):
    # Save plot to memory buffer instead of disk
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf

# --- Main Logic ---
def generate_report(df, filename, passing_mark):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Detect Mode: Standard Class (S1-S3) or Subject Group (S4-S5)
    is_subject_group = 'Subject Block' in df.columns
    
    # 1. Prepare Data Structure
    report_units = [] # List of (Header_Text, Subject_Name, Data_Subset)
    
    # Identify Subjects (Columns ending in [Attitude])
    attitude_cols = [c for c in df.columns if str(c).endswith(' [Attitude]')]
    subjects = [c.replace(' [Attitude]', '') for c in attitude_cols]
    subjects = [s for s in subjects if s in df.columns]

    if is_subject_group:
        # Logic for S4-S5 Grouping
        df['Level'] = df['Subject Block'].apply(lambda x: str(x)[0] if pd.notna(x) else '')
        df = df[df['Level'].isin(['4', '5'])] # Filter valid levels
        df['Level_Name'] = 'S' + df['Level']
        
        levels = sorted(df['Level_Name'].unique())
        for lvl in levels:
            for sub in subjects:
                # Create a filtered dataframe for this unit
                sub_df = df[df['Level_Name'] == lvl].copy()
                cols = ['Subject Block', sub, f"{sub} [Attitude]"]
                # Standardize columns
                unit_data = sub_df[cols].copy()
                unit_data.columns = ['Group', 'Mark', 'Attitude']
                report_units.append({
                    'title': f"{lvl} {sub}", 
                    'subject': sub,
                    'df': unit_data,
                    'level': lvl
                })
    else:
        # Logic for S1-S3 Standard Class
        # Guess level from filename or default to generic
        level_match = re.search(r'(S[1-6])', filename, re.IGNORECASE)
        level_prefix = level_match.group(0).upper() if level_match else "Subject"

        for sub in subjects:
            cols = ['Class', sub, f"{sub} [Attitude]"]
            unit_data = df[cols].copy()
            unit_data.columns = ['Group', 'Mark', 'Attitude']
            report_units.append({
                'title': f"{level_prefix} {sub}",
                'subject': sub,
                'df': unit_data,
                'level': level_prefix
            })

    # 2. Loop through units and build report
    for i, unit in enumerate(report_units):
        sub_df = unit['df']
        subject_name = unit['subject']
        report_title = unit['title']
        level = unit['level']
        
        # Clean Data
        sub_df['Mark'] = sub_df['Mark'].apply(clean_numeric)
        sub_df = sub_df.dropna(subset=['Group'])
        
        # Check emptiness
        if sub_df['Mark'].isna().all() and sub_df['Attitude'].isna().all():
            continue # Skip empty subjects

        # Determine Report Type (Restricted vs Full)
        # Rule: "Enhancement Class" subjects (except Spanish) only show Part 1
        is_enhancement = "Enhancement Class" in subject_name
        is_spanish = "Spanish" in subject_name
        is_restricted = is_enhancement and not is_spanish

        # Get Sorted Groups (Classes or Blocks)
        sorted_groups = get_sorted_groups(sub_df['Group'].unique())

        # Start New Section in Word
        if i == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
        
        # Header
        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs: p._element.getparent().remove(p._element)
        header_para = header.add_paragraph()
        header_para.text = report_title
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Title
        doc.add_heading(f"2526 Term 1 {report_title} Analysis Report", level=1)

        # --- Part 1: Attitude ---
        att_df = sub_df[sub_df['Attitude'].isin(['A', 'B', 'C', 'D'])]
        att_counts = att_df.groupby(['Group', 'Attitude']).size().unstack(fill_value=0)
        for g in ['A', 'B', 'C', 'D']:
             if g not in att_counts.columns: att_counts[g] = 0
        att_counts = att_counts[['A', 'B', 'C', 'D']] # Enforce order
        att_counts.loc['Whole Form'] = att_counts.sum() # Add Whole Form
        
        # Reindex
        plot_order = sorted_groups + ['Whole Form']
        att_counts = att_counts.reindex(plot_order).fillna(0)
        
        # Calculate %
        row_sums = att_counts.sum(axis=1)
        att_pct = att_counts.div(row_sums, axis=0) * 100
        att_pct = att_pct.fillna(0)

        # Plot
        fig, ax = plt.subplots(figsize=(9, 5))
        colors = sns.color_palette("Pastel1", 4)
        att_pct.plot(kind='bar', stacked=True, color=colors, width=0.7, edgecolor='grey', ax=ax)
        ax.set_title(f'Attitude to Learning - {subject_name}')
        ax.set_ylabel('Percentage')
        ax.set_xlabel('Class/Group')
        ax.legend(title='Attitude', bbox_to_anchor=(1.05, 1), loc='upper left')
        ax.set_ylim(0, 100)
        plt.tight_layout()
        
        doc.add_heading("1. Attitude to Learning", level=2)
        doc.add_paragraph("The chart below shows the distribution of 'Attitude to Learning' grades.")
        doc.add_picture(create_chart_image(fig), width=Inches(6))

        if is_restricted:
            continue # Stop here for Enhancement classes

        doc.add_page_break()

        # --- Part 2: Performance Stats ---
        stats_list = []
        stats_names = ['Mean', 'Standard Deviation', 'Maximum', 'Q3', 'Q2 (Median)', 'Q1', 'Minimum']
        
        # Calculate for groups
        for grp in sorted_groups:
            marks = sub_df[sub_df['Group'] == grp]['Mark'].dropna()
            if len(marks) > 0:
                desc = marks.describe()
                stats_list.append({
                    'Group': grp, 'Mean': desc['mean'], 'Standard Deviation': desc['std'],
                    'Maximum': desc['max'], 'Q3': desc['75%'], 'Q2 (Median)': desc['50%'],
                    'Q1': desc['25%'], 'Minimum': desc['min']
                })
            else:
                stats_list.append({'Group': grp})
        
        # Whole Form
        all_marks = sub_df['Mark'].dropna()
        if len(all_marks) > 0:
            desc = all_marks.describe()
            stats_list.append({
                'Group': 'Whole Form', 'Mean': desc['mean'], 'Standard Deviation': desc['std'],
                'Maximum': desc['max'], 'Q3': desc['75%'], 'Q2 (Median)': desc['50%'],
                'Q1': desc['25%'], 'Minimum': desc['min']
            })
            
        stats_df = pd.DataFrame(stats_list).set_index('Group').T.reindex(stats_names)

        doc.add_heading("2. Performance Statistics", level=2)
        doc.add_paragraph("Key performance metrics.")
        
        table = doc.add_table(rows=stats_df.shape[0]+1, cols=stats_df.shape[1]+1)
        table.style = 'Table Grid'
        row = table.rows[0]
        row.cells[0].text = ""
        for j, col_name in enumerate(stats_df.columns):
            row.cells[j+1].text = str(col_name)
        for r, (idx, data_row) in enumerate(stats_df.iterrows()):
            row_cells = table.rows[r+1].cells
            row_cells[0].text = str(idx)
            for c, val in enumerate(data_row):
                row_cells[c+1].text = f"{val:.1f}" if pd.notna(val) else "-"

        # --- Part 3: Boxplot ---
        fig, ax = plt.subplots(figsize=(10, 6))
        
        # 1. Create a clean DataFrame with just Group and Mark
        plot_df = sub_df[['Group', 'Mark']].dropna().copy()
        
        # 2. Create a duplicate for the 'Whole Form'
        whole_form_df = plot_df.copy()
        whole_form_df['Group'] = 'Whole Form'
        
        # 3. Combine them together
        combined_df = pd.concat([plot_df, whole_form_df])
        
        # 4. Enforce the exact sorting order for the X-axis
        boxplot_labels = sorted_groups + ['Whole Form']
        combined_df['Group'] = pd.Categorical(combined_df['Group'], categories=boxplot_labels, ordered=True)
        
        # 5. Plot using explicit x and y mappings
        sns.boxplot(data=combined_df, x='Group', y='Mark', palette="Pastel1", ax=ax)
        
        ax.set_xticklabels(boxplot_labels, rotation=45 if is_subject_group else 0)
        ax.set_title(f'Overall Mark Distribution - {subject_name}')
        ax.set_ylabel('Mark')
        ax.set_xlabel('') # Clear x-axis label to keep it neat
        ax.grid(axis='y', linestyle='--', alpha=0.7)
        plt.tight_layout()

        # --- THE MISSING LINES: Add the chart to the Word Document ---
        doc.add_heading("3. Overall Mark Distribution", level=2)
        doc.add_paragraph("Boxplot of Subject Mark distribution.")
        doc.add_picture(create_chart_image(fig), width=Inches(6))

# --- Part 4: Mark Dist Table ---
        bins = [0, 10, 20, 30, 40, 50, 60, 70, 80, 90, 101]
        labels = [f"{i}-{i+10} (excluding {i+10})" for i in range(0, 100, 10)] # Changed 90 to 100
        labels[-1] = "90-100"
        
        sub_df['Range'] = pd.cut(sub_df['Mark'], bins=bins, labels=labels, right=False)
        dist_counts = sub_df.groupby(['Group', 'Range'], observed=False).size().unstack(fill_value=0)
        
        # Ensure all groups exist
        for grp in sorted_groups:
             if grp not in dist_counts.index: dist_counts.loc[grp] = 0
        dist_counts.loc['Whole Form'] = dist_counts.sum()
        
        dist_table = dist_counts.T[sorted_groups + ['Whole Form']]
        
        # Stats
        attended, passed, rates = [], [], []
        for col in dist_table.columns:
            if col == 'Whole Form': m = sub_df['Mark'].dropna()
            else: m = sub_df[sub_df['Group'] == col]['Mark'].dropna()
            
            att_count = len(m)
            pass_count = len(m[m >= passing_mark])
            rate = (pass_count / att_count * 100) if att_count > 0 else 0
            
            attended.append(att_count)
            passed.append(pass_count)
            rates.append(f"{rate:.1f}%")

        dist_table.loc['Number of students attended'] = attended
        dist_table.loc['Number of students passed'] = passed
        dist_table.loc['Passing Rate'] = rates

        doc.add_heading("4. Mark Distribution Table", level=2)
        doc.add_paragraph("Frequency distribution of marks.")
        
        table = doc.add_table(rows=dist_table.shape[0]+1, cols=dist_table.shape[1]+1)
        table.style = 'Table Grid'
        row = table.rows[0]
        row.cells[0].text = ""
        for j, col_name in enumerate(dist_table.columns):
            row.cells[j+1].text = str(col_name)
        for r, (idx, data_row) in enumerate(dist_table.iterrows()):
            row_cells = table.rows[r+1].cells
            row_cells[0].text = str(idx)
            for c, val in enumerate(data_row):
                row_cells[c+1].text = str(val)

        doc.add_page_break()

        # --- Part 5: Grade Dist Chart ---
        fig, ax = plt.subplots(figsize=(10, 6))
        plot_data = dist_table.iloc[:-3].copy() # Exclude stats
        plot_data.plot(kind='bar', width=0.8, colormap='Pastel1', edgecolor='black', alpha=0.8, ax=ax)
        ax.set_title(f'Mark Distribution - {subject_name}')
        ax.set_ylabel('Number of Students')
        ax.set_xlabel('Mark Range')
        ax.legend(title='Class/Group')
        plt.tight_layout()

        doc.add_heading("5. Mark Distribution Chart", level=2)
        doc.add_paragraph("Chart presenting Mark Distribution.")
        doc.add_picture(create_chart_image(fig), width=Inches(6))

    # Save to buffer
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- Web Interface ---
st.title("📊 School Analytical Report Generator")
st.markdown("Upload your **Subject CSV** or **Excel** file to generate the MS Word report.")

uploaded_file = st.file_uploader("Upload Data File", type=['csv', 'xlsx'])
passing_mark = st.number_input("Passing Mark", min_value=0, max_value=100, value=40, step=1)

if uploaded_file is not None:
    st.info(f"File uploaded: {uploaded_file.name}")
    
    if st.button("Generate Report"):
        with st.spinner('Processing data and generating charts...'):
            try:
               # Load Data
                if uploaded_file.name.endswith('.csv'):
                    # Try reading with default UTF-8 first
                    try:
                        df_temp = pd.read_csv(uploaded_file, sep='\t', encoding='utf-8')
                    except UnicodeDecodeError:
                        # If it fails, fallback to UTF-16 (Common for Excel 'Unicode Text' exports)
                        uploaded_file.seek(0)
                        df_temp = pd.read_csv(uploaded_file, sep='\t', encoding='utf-16')

                    # Check if we need to skip the first row (header=1)
                    if 'Class' not in df_temp.columns and 'Subject Block' not in df_temp.columns:
                         uploaded_file.seek(0)
                         try:
                             df = pd.read_csv(uploaded_file, sep='\t', header=1, encoding='utf-8')
                         except UnicodeDecodeError:
                             uploaded_file.seek(0)
                             df = pd.read_csv(uploaded_file, sep='\t', header=1, encoding='utf-16')
                    else:
                         df = df_temp
                else:
                    df = pd.read_excel(uploaded_file)
                # Generate
                doc_file = generate_report(df, uploaded_file.name, passing_mark)
                
                # Success & Download
                st.success("Report generated successfully!")
                
                output_name = f"Analysis_Report_{uploaded_file.name.split('.')[0]}.docx"
                
                st.download_button(
                    label="Download MS Word Report",
                    data=doc_file,
                    file_name=output_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            except Exception as e:
                st.error(f"An error occurred: {e}")
                st.write("Please ensure the file format matches the school template (Tab-separated CSV or Excel).")
